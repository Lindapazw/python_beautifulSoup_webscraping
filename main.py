from bs4 import BeautifulSoup
import requests
from docx import Document

# website
website = 'https://subslikescript.com/movie/Titanic-120338'

# pongo en la variable result el resultado
result = requests.get(website)

# creo la variable content y coloco el result y obtengo el texto
content = result.text

# creamos la variable soup para localizar elementos dentro de la web
soup = BeautifulSoup(content, 'lxml')

# creamos un box con el contenido del article
box = soup.find('article', class_='main-article')

# en lugar de soup.find ponemos box para referirnos al article
title = box.find('h1').get_text()

# creamos la variable transcript
transcript = box.find('div', class_='full-script').get_text(strip=True, separator=' ')

# imprimimos
print(title)
print(transcript)

# Creamos un archivo de Word
document = Document()

# Agregamos el título al archivo de Word
document.add_heading(title, 0)

# Agregamos el texto al archivo de Word
document.add_paragraph(transcript)

# Guardamos el archivo
document.save('transcript.docx')